"""ChromaDB Service - Lightweight Vector Database for Knowledge Base.

This service provides:
- Fast vector similarity search
- Simple Python integration (no external services needed)
- LangChain/LlamaIndex compatible
- Auto-fallback when RAGFlow is unavailable

Usage:
    from app.services.chromadb_service import get_chromadb_service
    
    service = get_chromadb_service()
    await service.add_documents(agent_id, documents)
    results = await service.search(agent_id, query, top_k=5)
"""

import asyncio
from pathlib import Path
from typing import Any

from ..core.logger import get_logger

logger = get_logger(__name__)

# Lazy import to avoid startup overhead
_chromadb_client = None
_chromadb_imported = False


def _ensure_chromadb():
    """Lazy import ChromaDB."""
    global _chromadb_imported
    if not _chromadb_imported:
        try:
            import chromadb  # noqa: F401
            _chromadb_imported = True
        except ImportError:
            logger.warning("ChromaDB not installed. Run: pip install chromadb")
            raise ImportError("ChromaDB not installed")


class ChromaDBService:
    """ChromaDB service for vector storage and search."""
    
    def __init__(self, persist_directory: str = "/app/data/assets/chromadb"):
        """Initialize ChromaDB with persistent storage.
        
        Args:
            persist_directory: Directory to store ChromaDB data
        """
        self.persist_directory = persist_directory
        self._client = None
        self._collections: dict[str, Any] = {}
    
    @property
    def client(self):
        """Get or create ChromaDB client (lazy initialization)."""
        if self._client is None:
            _ensure_chromadb()
            import chromadb
            from chromadb.config import Settings
            
            # Ensure directory exists
            Path(self.persist_directory).mkdir(parents=True, exist_ok=True)
            
            self._client = chromadb.PersistentClient(
                path=self.persist_directory,
                settings=Settings(
                    anonymized_telemetry=False,
                    allow_reset=True,
                )
            )
            logger.info(f"ChromaDB initialized at {self.persist_directory}")
        
        return self._client
    
    def _get_collection_name(self, agent_id: str) -> str:
        """Generate collection name for an agent."""
        # ChromaDB collection names: 3-63 chars, alphanumeric + underscore
        safe_id = agent_id.replace("-", "_")[:50]
        return f"kb_{safe_id}"
    
    def get_collection(self, agent_id: str):
        """Get or create collection for an agent."""
        collection_name = self._get_collection_name(agent_id)
        
        if collection_name not in self._collections:
            self._collections[collection_name] = self.client.get_or_create_collection(
                name=collection_name,
                metadata={"agent_id": agent_id}
            )
        
        return self._collections[collection_name]
    
    async def add_documents(
        self,
        agent_id: str,
        documents: list[str],
        metadatas: list[dict] | None = None,
        ids: list[str] | None = None,
    ) -> dict:
        """Add documents to agent's knowledge base.
        
        Args:
            agent_id: Agent identifier
            documents: List of text documents
            metadatas: Optional list of metadata dicts
            ids: Optional list of unique IDs
            
        Returns:
            Dict with count of added documents
        """
        if not documents:
            return {"added": 0}
        
        collection = self.get_collection(agent_id)
        
        # Generate IDs if not provided
        if not ids:
            import uuid
            ids = [str(uuid.uuid4()) for _ in documents]
        
        # Default metadata
        if not metadatas:
            metadatas = [{"source": "manual"}] * len(documents)
        
        # Run in thread pool to avoid blocking
        loop = asyncio.get_event_loop()
        await loop.run_in_executor(
            None,
            lambda: collection.add(
                documents=documents,
                metadatas=metadatas,
                ids=ids,
            )
        )
        
        logger.info(f"Added {len(documents)} documents to agent {agent_id}")
        return {"added": len(documents), "ids": ids}
    
    async def add_text(
        self,
        agent_id: str,
        title: str,
        content: str,
        image_url: str | None = None,
        video_url: str | None = None,
    ) -> dict:
        """Add a single text document with title and optional media URLs.
        
        Args:
            agent_id: Agent identifier
            title: Document title
            content: Text content
            image_url: Optional image URL for device display
            video_url: Optional video URL for device display
            
        Returns:
            Dict with document ID
        """
        import uuid
        doc_id = str(uuid.uuid4())
        
        metadata = {
            "title": title, 
            "source": "text", 
            "type": "manual",
        }
        # Only include non-None/non-empty URLs in metadata
        if image_url and image_url.strip():
            metadata["image_url"] = image_url.strip()
        if video_url and video_url.strip():
            metadata["video_url"] = video_url.strip()
        
        await self.add_documents(
            agent_id=agent_id,
            documents=[f"{title}\n\n{content}"],
            metadatas=[metadata],
            ids=[doc_id],
        )
        
        return {"id": doc_id, "title": title, "image_url": image_url, "video_url": video_url}
    
    async def add_document(
        self,
        collection_id: str,
        doc_id: str,
        content: str,
        metadata: dict | None = None,
    ) -> dict:
        """Add a single document to knowledge base.
        
        API compatibility method for knowledge_base_entries.py
        
        Args:
            collection_id: Knowledge base ID (used as agent_id)
            doc_id: Unique document ID
            content: Document content
            metadata: Optional metadata dict
            
        Returns:
            Dict with document ID
        """
        from datetime import datetime
        
        # Use collection_id as agent_id
        agent_id = collection_id
        
        # Build metadata
        meta = {
            "doc_id": doc_id,
            "source": metadata.get("source", "manual") if metadata else "manual",
            "doc_type": metadata.get("doc_type", "text") if metadata else "text",
            "title": metadata.get("title", "") if metadata else "",
            "image_url": metadata.get("image_url", "") if metadata else "",
            "created_at": metadata.get("created_at", datetime.utcnow().isoformat()) if metadata else datetime.utcnow().isoformat(),
            "updated_at": metadata.get("updated_at", datetime.utcnow().isoformat()) if metadata else datetime.utcnow().isoformat(),
        }
        
        await self.add_documents(
            agent_id=agent_id,
            documents=[content],
            metadatas=[meta],
            ids=[doc_id],
        )
        
        return {"id": doc_id, "success": True}
    
    async def add_from_pdf(
        self,
        agent_id: str,
        file_content: bytes,
        filename: str,
    ) -> dict:
        """Import data from PDF file.
        
        Extracts text from all pages and stores in chunks.
        
        Args:
            agent_id: Agent identifier
            file_content: PDF file bytes
            filename: Original filename
            
        Returns:
            Dict with import results
        """
        import io
        
        # Try PyMuPDF (fitz) first, fallback to pypdf
        text_chunks = []
        
        try:
            import fitz  # PyMuPDF
            
            doc = fitz.open(stream=file_content, filetype="pdf")
            
            for page_num, page in enumerate(doc, start=1):
                text = page.get_text()
                if text.strip():
                    text_chunks.append({
                        "content": text,
                        "page": page_num,
                    })
            
            doc.close()
            logger.info(f"Extracted {len(text_chunks)} pages from PDF using PyMuPDF")
            
        except ImportError:
            # Fallback to pypdf
            try:
                from pypdf import PdfReader
                
                pdf_reader = PdfReader(io.BytesIO(file_content))
                
                for page_num, page in enumerate(pdf_reader.pages, start=1):
                    text = page.extract_text()
                    if text.strip():
                        text_chunks.append({
                            "content": text,
                            "page": page_num,
                        })
                
                logger.info(f"Extracted {len(text_chunks)} pages from PDF using pypdf")
                
            except ImportError:
                logger.error("No PDF library available. Install 'pymupdf' or 'pypdf'")
                raise ValueError("PDF extraction not available. Please install 'pip install pymupdf' or 'pip install pypdf'")
        
        if not text_chunks:
            return {"added": 0, "message": "No text found in PDF"}
        
        # Split into smaller chunks if pages are too long
        documents = []
        metadatas = []
        
        max_chunk_size = 4000  # Characters per chunk
        
        for chunk in text_chunks:
            content = chunk["content"]
            page = chunk["page"]
            
            # Split long pages into smaller chunks
            if len(content) > max_chunk_size:
                for i in range(0, len(content), max_chunk_size):
                    sub_content = content[i:i + max_chunk_size]
                    documents.append(sub_content)
                    metadatas.append({
                        "title": f"{filename} - Page {page}",
                        "source": filename,
                        "type": "pdf",
                        "page": page,
                        "chunk": i // max_chunk_size + 1,
                    })
            else:
                documents.append(content)
                metadatas.append({
                    "title": f"{filename} - Page {page}",
                    "source": filename,
                    "type": "pdf",
                    "page": page,
                })
        
        await self.add_documents(
            agent_id=agent_id,
            documents=documents,
            metadatas=metadatas,
        )
        
        return {
            "added": len(documents),
            "source": filename,
            "pages": len(text_chunks),
        }
    
    async def add_from_docx(
        self,
        agent_id: str,
        file_content: bytes,
        filename: str,
    ) -> dict:
        """Import data from Word DOCX file.
        
        Args:
            agent_id: Agent identifier
            file_content: DOCX file bytes
            filename: Original filename
            
        Returns:
            Dict with import results
        """
        import io
        
        try:
            from docx import Document
        except ImportError:
            raise ValueError("DOCX extraction not available. Please install 'pip install python-docx'")
        
        doc = Document(io.BytesIO(file_content))
        
        # Extract paragraphs
        paragraphs = []
        current_section = []
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                if current_section:
                    paragraphs.append("\n".join(current_section))
                    current_section = []
            else:
                current_section.append(text)
        
        if current_section:
            paragraphs.append("\n".join(current_section))
        
        # Also extract tables
        for table in doc.tables:
            table_text = []
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                table_text.append(row_text)
            if table_text:
                paragraphs.append("\n".join(table_text))
        
        if not paragraphs:
            return {"added": 0, "message": "No text found in DOCX"}
        
        # Combine into chunks
        documents = []
        metadatas = []
        
        max_chunk_size = 4000
        current_chunk = ""
        chunk_num = 1
        
        for para in paragraphs:
            if len(current_chunk) + len(para) > max_chunk_size:
                if current_chunk:
                    documents.append(current_chunk)
                    metadatas.append({
                        "title": f"{filename} - Section {chunk_num}",
                        "source": filename,
                        "type": "docx",
                        "section": chunk_num,
                    })
                    chunk_num += 1
                current_chunk = para
            else:
                current_chunk = f"{current_chunk}\n\n{para}" if current_chunk else para
        
        if current_chunk:
            documents.append(current_chunk)
            metadatas.append({
                "title": f"{filename} - Section {chunk_num}",
                "source": filename,
                "type": "docx",
                "section": chunk_num,
            })
        
        await self.add_documents(
            agent_id=agent_id,
            documents=documents,
            metadatas=metadatas,
        )
        
        return {
            "added": len(documents),
            "source": filename,
            "sections": chunk_num,
        }
    
    async def add_from_excel(
        self,
        agent_id: str,
        file_content: bytes,
        filename: str,
    ) -> dict:
        """Import data from Excel file.
        
        Expects Excel with columns: title, content (or just content)
        
        Args:
            agent_id: Agent identifier
            file_content: Excel file bytes
            filename: Original filename
            
        Returns:
            Dict with import results
        """
        import io
        import openpyxl
        
        # Load workbook
        wb = openpyxl.load_workbook(io.BytesIO(file_content), read_only=True)
        ws = wb.active
        
        # Get headers
        headers = [cell.value.lower() if cell.value else "" for cell in ws[1]]
        
        # Find content column
        content_col = None
        title_col = None
        
        for i, h in enumerate(headers):
            if "content" in h or "nội dung" in h:
                content_col = i
            if "title" in h or "tiêu đề" in h:
                title_col = i
        
        if content_col is None:
            # Use first column as content
            content_col = 0
        
        # Extract data
        documents = []
        metadatas = []
        
        for row_num, row in enumerate(ws.iter_rows(min_row=2, values_only=True), start=2):
            content = row[content_col] if content_col < len(row) else None
            if not content:
                continue
                
            title = row[title_col] if title_col is not None and title_col < len(row) else f"Row {row_num}"
            
            documents.append(f"{title}\n\n{content}")
            metadatas.append({
                "title": str(title),
                "source": filename,
                "type": "excel",
                "row": row_num,
            })
        
        wb.close()
        
        if not documents:
            return {"added": 0, "message": "No data found in Excel"}
        
        await self.add_documents(
            agent_id=agent_id,
            documents=documents,
            metadatas=metadatas,
        )
        
        return {"added": len(documents), "source": filename}
    
    async def add_from_google_sheets(
        self,
        agent_id: str,
        sheet_url: str,
    ) -> dict:
        """Import data from Google Sheets.
        
        Sheet must be public or shared with "Anyone with link".
        Expects columns: title, content (or just content)
        
        Args:
            agent_id: Agent identifier
            sheet_url: Google Sheets URL
            
        Returns:
            Dict with import results
        """
        import re
        import httpx
        
        # Extract sheet ID from URL
        match = re.search(r'/d/([a-zA-Z0-9-_]+)', sheet_url)
        if not match:
            raise ValueError("Invalid Google Sheets URL")
        
        sheet_id = match.group(1)
        
        # Use CSV export (works for public sheets)
        csv_url = f"https://docs.google.com/spreadsheets/d/{sheet_id}/export?format=csv"
        
        async with httpx.AsyncClient() as client:
            response = await client.get(csv_url, follow_redirects=True)
            if response.status_code != 200:
                raise ValueError(f"Cannot access Google Sheet. Make sure it's public. Status: {response.status_code}")
            
            csv_content = response.text
        
        # Parse CSV
        import csv
        import io
        
        reader = csv.DictReader(io.StringIO(csv_content))
        rows = list(reader)
        
        if not rows:
            return {"added": 0, "message": "Sheet is empty"}
        
        # Find content column - support multiple formats
        headers = list(rows[0].keys())
        content_col = None
        title_col = None
        id_col = None
        image_col = None  # NEW: Image URL column
        
        for h in headers:
            h_lower = h.lower().strip()
            # Content/Detail column
            if h_lower in ["content", "nội dung", "detail", "chi tiết", "mô tả", "description"]:
                content_col = h
            # Title/Name column  
            if h_lower in ["title", "tiêu đề", "name", "tên", "tên sản phẩm"]:
                title_col = h
            # ID column
            if h_lower in ["id", "mã", "code", "mã sp", "sku"]:
                id_col = h
            # Image URL column - NEW
            if h_lower in ["image_url", "image", "hình", "hình ảnh", "url_image", "img", "photo", "ảnh"]:
                image_col = h
        
        # Fallback: if only 3-4 columns and no matches, assume id/name/detail/image format
        if content_col is None and len(headers) >= 3:
            id_col = headers[0]
            title_col = headers[1]
            content_col = headers[2]
            if len(headers) >= 4:
                image_col = headers[3]
        elif content_col is None:
            content_col = headers[-1]  # Use last column as content
        
        # Extract data
        documents = []
        metadatas = []
        
        for i, row in enumerate(rows, start=2):
            content = row.get(content_col, "")
            if not content:
                continue
            
            title = row.get(title_col, f"Row {i}") if title_col else f"Row {i}"
            doc_id = row.get(id_col, "") if id_col else ""
            image_url = row.get(image_col, "") if image_col else ""  # NEW
            
            # Format document: include id if available
            if doc_id:
                doc_text = f"[{doc_id}] {title}\n\n{content}"
            else:
                doc_text = f"{title}\n\n{content}"
            
            documents.append(doc_text)
            metadatas.append({
                "title": str(title),
                "source": f"google_sheets:{sheet_id}",
                "type": "google_sheets",
                "row": i,
                "original_id": str(doc_id) if doc_id else None,
                "image_url": str(image_url).strip() if image_url else None,  # NEW
            })
        
        if not documents:
            return {"added": 0, "message": "No data found in Sheet"}
        
        await self.add_documents(
            agent_id=agent_id,
            documents=documents,
            metadatas=metadatas,
        )
        
        return {"added": len(documents), "source": sheet_url}
    
    async def search(
        self,
        agent_id: str,
        query: str,
        top_k: int = 5,
    ) -> list[dict]:
        """Search for similar documents.
        
        Args:
            agent_id: Agent identifier
            query: Search query
            top_k: Number of results to return
            
        Returns:
            List of matching documents with scores
        """
        collection = self.get_collection(agent_id)
        
        # Run in thread pool
        loop = asyncio.get_event_loop()
        results = await loop.run_in_executor(
            None,
            lambda: collection.query(
                query_texts=[query],
                n_results=top_k,
                include=["documents", "metadatas", "distances"],
            )
        )
        
        # Format results
        chunks = []
        documents = results.get("documents", [[]])[0]
        metadatas = results.get("metadatas", [[]])[0]
        distances = results.get("distances", [[]])[0]
        
        for doc, meta, dist in zip(documents, metadatas, distances):
            # Convert distance to similarity score (0-1)
            # ChromaDB uses L2 distance by default
            score = max(0, 1 - dist / 2)  # Approximate normalization
            
            chunks.append({
                "content": doc,
                "source": meta.get("source", "unknown"),
                "title": meta.get("title", ""),
                "score": score,
                "image_url": meta.get("image_url"),  # Image URL for device display
                "video_url": meta.get("video_url"),  # Video URL for device display
                "metadata": meta,
            })
        
        return chunks
    
    async def delete_collection(self, agent_id: str) -> bool:
        """Delete agent's knowledge base collection.
        
        Args:
            agent_id: Agent identifier
            
        Returns:
            True if deleted
        """
        collection_name = self._get_collection_name(agent_id)
        
        try:
            self.client.delete_collection(collection_name)
            self._collections.pop(collection_name, None)
            logger.info(f"Deleted collection for agent {agent_id}")
            return True
        except Exception as e:
            logger.warning(f"Failed to delete collection: {e}")
            return False
    
    async def get_stats(self, agent_id: str) -> dict:
        """Get statistics for agent's knowledge base.
        
        Args:
            agent_id: Agent identifier
            
        Returns:
            Dict with collection stats
        """
        collection = self.get_collection(agent_id)
        count = collection.count()
        
        return {
            "agent_id": agent_id,
            "document_count": count,
            "collection_name": self._get_collection_name(agent_id),
        }
    
    async def list_documents(
        self,
        agent_id: str,
        limit: int = 100,
        offset: int = 0,
    ) -> dict:
        """List all documents in agent's knowledge base.
        
        Args:
            agent_id: Agent identifier
            limit: Max documents to return
            offset: Pagination offset
            
        Returns:
            Dict with documents list
        """
        collection = self.get_collection(agent_id)
        
        # Get all documents
        loop = asyncio.get_event_loop()
        results = await loop.run_in_executor(
            None,
            lambda: collection.get(
                include=["documents", "metadatas"],
                limit=limit,
                offset=offset,
            )
        )
        
        documents = []
        ids = results.get("ids", [])
        docs = results.get("documents", [])
        metas = results.get("metadatas", [])
        
        for doc_id, doc, meta in zip(ids, docs, metas):
            documents.append({
                "id": doc_id,
                "content": doc[:200] + "..." if len(doc) > 200 else doc,
                "title": meta.get("title", ""),
                "source": meta.get("source", "manual"),
                "type": meta.get("type", "unknown"),
                "image_url": meta.get("image_url"),  # NEW
            })
        
        return {
            "documents": documents,
            "total": collection.count(),
        }
    
    async def delete_document(self, agent_id: str, doc_id: str) -> bool:
        """Delete a specific document.
        
        Args:
            agent_id: Agent identifier
            doc_id: Document ID to delete
            
        Returns:
            True if deleted
        """
        collection = self.get_collection(agent_id)
        
        loop = asyncio.get_event_loop()
        await loop.run_in_executor(
            None,
            lambda: collection.delete(ids=[doc_id])
        )
        
        logger.info(f"Deleted document {doc_id} from agent {agent_id}")
        return True
    
    async def update_document(
        self,
        agent_id: str,
        doc_id: str,
        content: str,
        title: str | None = None,
    ) -> bool:
        """Update a specific document/chunk.
        
        Args:
            agent_id: Agent identifier
            doc_id: Document ID to update
            content: New content
            title: New title (optional)
            
        Returns:
            True if updated
        """
        collection = self.get_collection(agent_id)
        
        # Get current metadata
        loop = asyncio.get_event_loop()
        current = await loop.run_in_executor(
            None,
            lambda: collection.get(ids=[doc_id], include=["metadatas"])
        )
        
        if not current.get("ids"):
            raise ValueError(f"Document {doc_id} not found")
        
        # Update metadata with new title if provided
        metadata = current.get("metadatas", [{}])[0] or {}
        if title:
            metadata["title"] = title
        
        # Update document
        await loop.run_in_executor(
            None,
            lambda: collection.update(
                ids=[doc_id],
                documents=[content],
                metadatas=[metadata],
            )
        )
        
        logger.info(f"Updated document {doc_id} in agent {agent_id}")
        return True
    
    async def health_check(self) -> dict:
        """Check ChromaDB health.
        
        Returns:
            Health status dict
        """
        try:
            _ensure_chromadb()
            # Try to access client
            _ = self.client.heartbeat()
            return {
                "status": "healthy",
                "service": "chromadb",
                "persist_directory": self.persist_directory,
            }
        except Exception as e:
            return {
                "status": "unhealthy",
                "service": "chromadb",
                "error": str(e),
            }


# Singleton instance
_chromadb_service: ChromaDBService | None = None


def get_chromadb_service() -> ChromaDBService:
    """Get singleton ChromaDB service instance."""
    global _chromadb_service
    if _chromadb_service is None:
        _chromadb_service = ChromaDBService()
    return _chromadb_service
